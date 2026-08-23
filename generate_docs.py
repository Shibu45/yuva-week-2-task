import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_workflow_diagram(output_path="workflow_flowchart.png"):
    fig, ax = plt.subplots(figsize=(10, 5.5), dpi=300)
    ax.axis('off')
    fig.patch.set_facecolor('#F8FAFC')
    
    colors = {
        'source': '#1E3A8A',     # Dark Navy
        'ingest': '#2563EB',     # Blue
        'quality': '#0D9488',    # Teal
        'clean': '#D97706',      # Amber
        'transform': '#7C3AED',  # Purple
        'export': '#059669'      # Emerald Green
    }
    
    boxes = [
        ("1. Data Sourcing", "Public APIs, Scraping,\nDB Connectors & CSVs", 0.04, 0.6, colors['source']),
        ("2. Ingestion & Schema", "Format Parsing, Type Check,\nPydantic Validation", 0.19, 0.6, colors['ingest']),
        ("3. Quality & Auditing", "Missingness Analysis,\n6-Dimension Checks", 0.34, 0.6, colors['quality']),
        ("4. Outliers & Cleaning", "IQR / Z-Score / IsoForest,\nKNN & MICE Imputation", 0.49, 0.6, colors['clean']),
        ("5. Feature Transform", "Robust Scaling, One-Hot/\nTarget Encoding & Log Tx", 0.64, 0.6, colors['transform']),
        ("6. Analytics Ready", "Parquet Export, Feature Store,\nDownstream Model Input", 0.79, 0.6, colors['export'])
    ]
    
    box_width = 0.14
    box_height = 0.28
    
    for title, desc, x, y, col in boxes:
        rect = patches.FancyBboxPatch(
            (x, y - box_height/2), box_width, box_height,
            boxstyle="round,pad=0.02,rounding_size=0.03",
            linewidth=1.5, edgecolor=col, facecolor='white',
            zorder=2
        )
        ax.add_patch(rect)
        
        header_rect = patches.FancyBboxPatch(
            (x, y + box_height/2 - 0.07), box_width, 0.07,
            boxstyle="square,pad=0",
            linewidth=0, facecolor=col,
            zorder=3
        )
        ax.add_patch(header_rect)
        
        ax.text(x + box_width/2, y + box_height/2 - 0.035, title,
                color='white', weight='bold', fontsize=8.5, ha='center', va='center', zorder=4)
        
        ax.text(x + box_width/2, y - 0.03, desc,
                color='#1E293B', fontsize=7.5, ha='center', va='center', zorder=4, multialignment='center')

    for i in range(len(boxes) - 1):
        x_start = boxes[i][2] + box_width
        x_end = boxes[i+1][2]
        y_pos = boxes[i][3]
        
        ax.annotate('', xy=(x_end, y_pos), xytext=(x_start, y_pos),
                    arrowprops=dict(arrowstyle="->", color="#64748B", lw=2, mutation_scale=15),
                    zorder=1)
        
    ax.annotate('', xy=(boxes[2][2] + box_width/2, boxes[2][3] - box_height/2 - 0.02),
                xytext=(boxes[4][2] + box_width/2, boxes[4][3] - box_height/2 - 0.02),
                arrowprops=dict(arrowstyle="->", color="#94A3B8", lw=1.5, ls='--',
                                connectionstyle="arc3,rad=0.3", mutation_scale=12),
                zorder=1)
    ax.text(0.53, 0.22, "Iterative Quality Audit & Re-tuning Loop",
            color='#64748B', fontsize=8, style='italic', ha='center', va='center')

    plt.title("End-to-End Data Acquisition & Preprocessing Strategy Pipeline",
              fontsize=13, weight='bold', color='#0F172A', pad=25)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"Workflow diagram saved to {output_path}")

def create_timeline_diagram(output_path="timeline_gantt.png"):
    fig, ax = plt.subplots(figsize=(10, 4.5), dpi=300)
    fig.patch.set_facecolor('#F8FAFC')
    ax.set_facecolor('#FFFFFF')
    
    phases = [
        "Phase 1: Sourcing & Requirements",
        "Phase 2: Ingestion & Validation",
        "Phase 3: Quality Audit & Cleaning",
        "Phase 4: Feature Transformation",
        "Phase 5: Integration & Export"
    ]
    
    start_hours = [0, 6, 12, 20, 28]
    durations = [6, 6, 8, 8, 7]
    colors = ['#1E3A8A', '#2563EB', '#0D9488', '#7C3AED', '#059669']
    
    y_pos = range(len(phases))
    
    for i in range(len(phases)):
        ax.barh(y_pos[i], durations[i], left=start_hours[i], height=0.5,
                color=colors[i], edgecolor='none', alpha=0.9, zorder=3)
        ax.text(start_hours[i] + durations[i]/2, y_pos[i], f"{durations[i]}h",
                color='white', weight='bold', fontsize=9, ha='center', va='center', zorder=4)
        
    ax.set_yticks(y_pos)
    ax.set_yticklabels(phases, fontsize=9.5, weight='bold', color='#1E293B')
    ax.invert_yaxis()
    
    ax.set_xlabel("Timeline (Hours 0 to 35)", fontsize=10, weight='bold', color='#0F172A', labelpad=10)
    ax.set_xlim(0, 35)
    ax.set_xticks(range(0, 36, 5))
    ax.grid(axis='x', linestyle='--', alpha=0.5, color='#CBD5E1', zorder=0)
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#94A3B8')
    ax.spines['bottom'].set_color('#94A3B8')
    
    plt.title("30-35 Hour Phased Data Preprocessing Execution Timeline",
              fontsize=12, weight='bold', color='#0F172A', pad=15)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"Timeline chart saved to {output_path}")

# XML Helper Functions for Styling Word Document
def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text, title="STRATEGIC IMPERATIVE"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "F0F9FF") # Light Blue Tint
    
    # Left border blue accent
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="single" w:sz="36" w:space="0" w:color="2563EB"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>
    '''
    tcPr.append(parse_xml(borders_xml))
    set_cell_margins(cell, top=140, bottom=140, left=200, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run_title = p.add_run(f"📌 {title}: ")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(37, 99, 235)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(30, 41, 59)
    
    # Empty paragraph after callout
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "F8FAFC") # Gray background
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
        <w:left w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
        <w:right w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
    </w:tcBorders>
    '''
    tcPr.append(parse_xml(borders_xml))
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)

def style_table_header(row, col_widths, bg_color="1E3A8A"):
    for idx, cell in enumerate(row.cells):
        cell.width = col_widths[idx]
        set_cell_shading(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

def style_table_rows(table, col_widths, stripe=True):
    for i, row in enumerate(table.rows[1:]):
        bg_color = "F1F5F9" if (i % 2 == 1 and stripe) else "FFFFFF"
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            # Subtle light gray bottom border
            tcPr = cell._tc.get_or_add_tcPr()
            borders_xml = f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:right w:val="none"/>
            </w:tcBorders>
            '''
            tcPr.append(parse_xml(borders_xml))
            
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(30, 41, 59)

def generate_doc(output_path="Data_Acquisition_and_Preprocessing_Strategy.docx"):
    doc = docx.Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Standard styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # Document Title Banner (Table)
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_cell = banner_table.cell(0, 0)
    banner_cell.width = Inches(6.5)
    set_cell_shading(banner_cell, "1E3A8A") # Navy Blue
    set_cell_margins(banner_cell, top=200, bottom=200, left=200, right=200)
    
    p = banner_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p.add_run("ENTERPRISE DATA ACQUISITION & PREPROCESSING STRATEGY")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = RGBColor(255, 255, 255)
    
    p_sub = banner_cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sub = p_sub.add_run("A Technical Roadmap for Data Sourcing, Quality Auditing, Outlier Detection, Imputation, and Feature Transformation in Python Analytics")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = RGBColor(226, 232, 240)
    
    # Meta Details
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(14)
    r_meta = p_meta.add_run("AUTHOR: Senior Data Lead / Analytics Engineer   |   VERSION: 1.0   |   TARGET TIMELINE: 30–35 Hours Execution")
    r_meta.font.size = Pt(8.5)
    r_meta.font.color.rgb = RGBColor(100, 116, 139)
    r_meta.bold = True

    # Callout Executive Summary
    add_callout_box(
        doc,
        "In modern data analytics, raw data is rarely directly consumable. Garbage in yields garbage out. "
        "This document presents a strategic framework for acquiring data from public APIs, databases, and web sources, "
        "enforcing 6-dimension data quality controls, addressing missingness (MCAR/MAR/MNAR) and anomalies (IQR/Z-Score/Isolation Forest), "
        "and performing optimal feature scaling and encoding in Python for maximum downstream analytical fidelity.",
        title="EXECUTIVE STRATEGY SUMMARY"
    )

    # Helper function for Headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(13, 148, 136) # Teal
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(124, 58, 237) # Purple
        return h

    # Section 1
    add_heading_1("1. Strategic Vision & Project Objectives")
    p = doc.add_paragraph(
        "Data acquisition and preprocessing form the foundational bedrock of any quantitative analytics endeavor. "
        "A rigorous preprocessing strategy ensures that real-world noise, structural irregularities, missing values, "
        "and distribution anomalies are systematically triaged before data enters exploratory, statistical, or machine learning pipelines."
    )
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Key Project Objectives:\n"
        "1. Standardize Multi-Source Ingestion: Integrate diverse external data streams (REST APIs, Web Scrapes, Database Extracts) seamlessly.\n"
        "2. Automate Quality Auditing: Implement a repeatable 6-dimension schema validation matrix.\n"
        "3. Mitigate Information Loss: Employ state-of-the-art imputation algorithms (MICE, KNN) customized to missingness types.\n"
        "4. Prevent Distortions: Apply robust anomaly detection (Isolation Forest, IQR) to safeguard analytical modeling.\n"
        "5. Optimize Analytical Fit: Tailor feature scaling and encoding to specific downstream algorithmic requirements."
    )
    p.paragraph_format.space_after = Pt(12)

    # Section 2
    add_heading_1("2. Public Data Sources & Selection Framework")
    p = doc.add_paragraph(
        "Selecting appropriate data sources is a multi-dimensional optimization task. To balance data richness, "
        "licensing constraints, and computational feasibility, public datasets are evaluated across a 7-point assessment matrix."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("2.1 Data Selection Criteria Matrix")
    
    # Table of Selection Criteria
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(1.8), Inches(1.5), Inches(3.2)]
    
    headers = ["Evaluation Criterion", "Threshold / Metric", "Strategic Rationale & Target Objective"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].paragraphs[0].text = text
    style_table_header(table.rows[0], widths)

    criteria_data = [
        ("1. Freshness & Cadence", "Daily / Real-time API", "Ensures models reflect current macro trends without temporal lag."),
        ("2. Licensing & Governance", "CC-BY 4.0 / Open Government", "Prevents legal liabilities and guarantees commercial/academic reuse."),
        ("3. Schema Stability", "Semantic Versioned API", "Minimizes breaking structural shifts during automated ingestion runs."),
        ("4. Missingness Ratio", "Unrecoverable Nulls < 15%", "Limits bias introduced by aggressive row/column truncation."),
        ("5. Spatial/Temporal Coverage", "5+ Years Longitudinal Data", "Provides sufficient sample size for statistical power and time-series trends."),
        ("6. Granularity", "Transaction / Event Level", "Enables flexible aggregation at daily, weekly, or regional levels."),
        ("7. API Throttling Limits", ">= 1,000 requests/min", "Guarantees high-throughput batch extraction within project schedule.")
    ]

    for row_idx, data in enumerate(criteria_data, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].paragraphs[0].text = data[0]
        row_cells[1].paragraphs[0].text = data[1]
        row_cells[2].paragraphs[0].text = data[2]

    style_table_rows(table, widths)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

    add_heading_2("2.2 Identified Public Benchmark Sources")
    p = doc.add_paragraph(
        "For this project, three benchmark public datasets have been prioritized:\n"
        "• World Bank Data API: Global macroeconomic metrics, development indicators, and trade volumes.\n"
        "• NYC OpenData Portal: High-density urban mobility, service requests, and transaction logs.\n"
        "• Kaggle Global E-Commerce & Financial Transactions Benchmark: Granular consumer behavior and pricing telemetry."
    )
    p.paragraph_format.space_after = Pt(12)

    # Section 3
    add_heading_1("3. Data Extraction Methods & Technical Challenge Mitigation")
    p = doc.add_paragraph(
        "Data extraction bridges public repositories and local analytics memory spaces. "
        "Each extraction pattern introduces distinct operational failure modes."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("3.1 Extraction Challenges & Technical Countermeasures")
    
    table_ext = doc.add_table(rows=6, cols=3)
    table_ext.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_ext = [Inches(1.6), Inches(2.2), Inches(2.7)]
    
    headers_ext = ["Extraction Vector", "Anticipated Challenge", "Technical Mitigation Strategy"]
    for idx, text in enumerate(headers_ext):
        table_ext.rows[0].cells[idx].paragraphs[0].text = text
    style_table_header(table_ext.rows[0], widths_ext)

    ext_data = [
        ("RESTful APIs", "HTTP 429 Rate Limiting & Throttling", "Exponential backoff retry policy (tenacity library) with persistent API key rotation."),
        ("Web Scraping", "Anti-Bot Captchas & Dynamic JS Rendering", "Headless Playwright rendering with user-agent randomization and polite delay intervals (1-3s)."),
        ("Bulk Databases", "Memory Overflow on High-Volume CSVs", "Chunked streaming with pandas read_csv(chunksize=50000) and PyArrow Parquet memory mapping."),
        ("Heterogeneous Data", "Character Encoding Mismatch (Latin-1/UTF-8)", "Chardet automatic encoding detection and forced UTF-8 byte stream coercion upon download."),
        ("Network Connections", "Socket Timeouts & Dropped Packets", "Atomic partial downloads stored in staging buckets with SHA-256 integrity checksum verification.")
    ]

    for row_idx, data in enumerate(ext_data, start=1):
        row_cells = table_ext.rows[row_idx].cells
        row_cells[0].paragraphs[0].text = data[0]
        row_cells[1].paragraphs[0].text = data[1]
        row_cells[2].paragraphs[0].text = data[2]

    style_table_rows(table_ext, widths_ext)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

    # Section 4
    add_heading_1("4. Data Quality Assurance, Validation & Cleaning Strategy")
    p = doc.add_paragraph(
        "Data quality auditing is executed across six foundational dimensions: Completeness, Accuracy, "
        "Consistency, Timeliness, Validity, and Uniqueness. Violations are isolated before downstream transformation."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("4.1 Missing Data Taxonomy & Imputation Framework")
    p = doc.add_paragraph(
        "Handling missing values requires diagnosing the statistical mechanism underlying the missingness:"
    )
    p.paragraph_format.space_after = Pt(4)

    add_callout_box(
        doc,
        "1. MCAR (Missing Completely at Random): Missingness is independent of observed and unobserved data. Action: Mean/Median Imputation or Listwise Deletion if < 3%.\n"
        "2. MAR (Missing at Random): Missingness depends on observed features (e.g., income missingness correlated with age). Action: KNN Imputation or MICE (Multivariate Imputation by Chained Equations).\n"
        "3. MNAR (Missing Not at Random): Missingness relates to unobserved true values. Action: Add binary missing indicator columns (is_missing=1) and domain-specific modeling.",
        title="MISSINGNESS MECHANISM TAXONOMY"
    )

    add_heading_2("4.2 Outlier Detection & Treatment Methodologies")
    p = doc.add_paragraph(
        "Outliers represent either legitimate extreme events or measurement errors. Three complementary techniques are utilized:"
    )
    p.paragraph_format.space_after = Pt(6)

    p_out = doc.add_paragraph(
        "• Z-Score (Parametric): Flags data points where |Z| > 3.0. Assumes underlying normal distribution.\n"
        "  Formula: Z = (X - μ) / σ\n\n"
        "• Interquartile Range (IQR - Non-Parametric): Resilient to existing heavy skewness. Flags points outside [Q1 - 1.5×IQR, Q3 + 1.5×IQR].\n"
        "  Formula: IQR = Q3 - Q1\n\n"
        "• Isolation Forest (Multivariate ML): Constructs isolation trees to partition high-dimensional feature spaces. Anomaly scores identify complex multivariate outliers."
    )
    p_out.paragraph_format.space_after = Pt(10)

    add_heading_2("4.3 Outlier Treatment Rationale")
    p = doc.add_paragraph(
        "Rather than blanket row deletion (which destroys sample size), extreme values undergo Winsorization "
        "(capping at 1st and 99th percentiles) or log transformation to preserve sample density while bounding leverage."
    )
    p.paragraph_format.space_after = Pt(12)

    # Section 5
    add_heading_1("5. Python Toolchain Justification & Transformation Rationale")
    p = doc.add_paragraph(
        "The selection of Python libraries is tailored to maximize execution efficiency, memory management, "
        "and algorithmic rigor."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("5.1 Technical Toolchain Rationale Matrix")

    table_tools = doc.add_table(rows=6, cols=3)
    table_tools.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_tools = [Inches(1.5), Inches(2.2), Inches(2.8)]
    
    headers_tools = ["Library / Framework", "Primary Preprocessing Functions", "Technical Justification & Advantages"]
    for idx, text in enumerate(headers_tools):
        table_tools.rows[0].cells[idx].paragraphs[0].text = text
    style_table_header(table_tools.rows[0], widths_tools)

    tools_data = [
        ("Pandas 2.x", "Dataframe operations, time-series alignment, group aggregations", "C-optimized Arrow backend support, intuitive index alignment, rich I/O connectors."),
        ("NumPy 1.26+", "Vectorized array math, matrix transformations, log transforms", "BLAS/LAPACK hardware acceleration, minimal memory overhead for numerical computations."),
        ("Scikit-Learn 1.4+", "StandardScaler, RobustScaler, OneHotEncoder, IterativeImputer", "Production-tested API, pipelining capabilities, prevention of data leakage via fit/transform splits."),
        ("SciPy 1.12+", "Shapiro-Wilk normality tests, Box-Cox transforms, Z-scores", "Extensive statistical distribution modeling and formal hypothesis testing utilities."),
        ("Polars / Dask", "Out-of-core evaluation, multi-threaded parallel file reads", "Used as high-scale fallback for datasets exceeding RAM limits (> 10GB).")
    ]

    for row_idx, data in enumerate(tools_data, start=1):
        row_cells = table_tools.rows[row_idx].cells
        row_cells[0].paragraphs[0].text = data[0]
        row_cells[1].paragraphs[0].text = data[1]
        row_cells[2].paragraphs[0].text = data[2]

    style_table_rows(table_tools, widths_tools)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

    add_heading_2("5.2 Data Transformation Methodology")
    p = doc.add_paragraph(
        "• Scaling Strategy: Continuous numeric features are transformed using RobustScaler (subtracts median, scales by IQR) "
        "when outliers exist, or StandardScaler (Z-scoring) for normally distributed metrics.\n"
        "• Categorical Encoding: Low-cardinality nominal values (< 10 levels) undergo One-Hot Encoding. High-cardinality nominals use Target / Frequency Encoding.\n"
        "• Skewness Correction: Features with skewness |S| > 1.0 are log-transformed (np.log1p) or Box-Cox transformed."
    )
    p.paragraph_format.space_after = Pt(12)

    # Section 6
    add_heading_1("6. Preprocessing Workflow Architecture & Pseudo-code")
    p = doc.add_paragraph(
        "The end-to-end data acquisition and preprocessing workflow is operationalized into an object-oriented Python pipeline architecture."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("6.1 Visual Pipeline Architecture Flowchart")
    if os.path.exists("workflow_flowchart.png"):
        doc.add_paragraph()
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture("workflow_flowchart.png", width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 1: End-to-End Modular Data Preprocessing Workflow Diagram")
        r_cap.font.size = Pt(8.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)

    add_heading_2("6.2 Modular Python Preprocessing Pseudo-code")
    code_snippet = """import pandas as pd
import numpy as np
from sklearn.base import BaseEstimator, TransformerMixin
from sklearn.preprocessing import RobustScaler, OneHotEncoder
from sklearn.impute import KNNImputer
from sklearn.ensemble import IsolationForest

class ProductionDataPipeline(BaseEstimator, TransformerMixin):
    def __init__(self, numeric_cols, categorical_cols, outlier_contamination=0.03):
        self.numeric_cols = numeric_cols
        self.categorical_cols = categorical_cols
        self.contamination = outlier_contamination
        
        self.imputer = KNNImputer(n_neighbors=5)
        self.scaler = RobustScaler()
        self.encoder = OneHotEncoder(sparse_output=False, handle_unknown='ignore')
        self.iso_forest = IsolationForest(contamination=self.contamination, random_state=42)
        
    def fit(self, X, y=None):
        # 1. Fit Missingness Imputer on Numerics
        self.imputer.fit(X[self.numeric_cols])
        imputed_num = self.imputer.transform(X[self.numeric_cols])
        
        # 2. Fit Robust Scaler
        self.scaler.fit(imputed_num)
        
        # 3. Fit Categorical Encoder
        self.encoder.fit(X[self.categorical_cols].fillna("MISSING"))
        
        # 4. Fit Anomaly Detector
        self.iso_forest.fit(imputed_num)
        return self

    def transform(self, X):
        X_clean = X.copy()
        
        # 1. Impute Numerics
        imputed_num = self.imputer.transform(X_clean[self.numeric_cols])
        
        # 2. Flag Outliers
        outlier_flags = self.iso_forest.predict(imputed_num)
        X_clean['is_outlier'] = np.where(outlier_flags == -1, 1, 0)
        
        # 3. Scale Numerics
        scaled_num = self.scaler.transform(imputed_num)
        df_scaled = pd.DataFrame(scaled_num, columns=[f"{c}_scaled" for c in self.numeric_cols], index=X_clean.index)
        
        # 4. Encode Categoricals
        encoded_cat = self.encoder.transform(X_clean[self.categorical_cols].fillna("MISSING"))
        cat_feature_names = self.encoder.get_feature_names_out(self.categorical_cols)
        df_encoded = pd.DataFrame(encoded_cat, columns=cat_feature_names, index=X_clean.index)
        
        # 5. Assemble Analytics-Ready Matrix
        final_df = pd.concat([df_scaled, df_encoded, X_clean[['is_outlier']]], axis=1)
        return final_df
"""
    add_code_block(doc, code_snippet)

    # Section 7
    add_heading_1("7. Coherence & Linkage to Downstream Analytics")
    p = doc.add_paragraph(
        "Every preprocessing decision directly influences the mathematical assumptions and performance of downstream analytical algorithms."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("7.1 Preprocessing Action to Analytics Outcome Matrix")

    table_link = doc.add_table(rows=6, cols=3)
    table_link.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_link = [Inches(1.8), Inches(1.8), Inches(2.9)]
    
    headers_link = ["Preprocessing Action", "Downstream Analytics Task", "Direct Analytical Outcome & Impact"]
    for idx, text in enumerate(headers_link):
        table_link.rows[0].cells[idx].paragraphs[0].text = text
    style_table_header(table_link.rows[0], widths_link)

    link_data = [
        ("IQR Winsorization", "Multiple Linear Regression", "Prevents high-leverage extreme values from skewing regression coefficients and inflating Standard Error."),
        ("Robust Scaling", "K-Means / KNN Clustering", "Ensures Euclidean distance metrics are not dominated by high-magnitude scales, equalizing feature weights."),
        ("MICE / KNN Imputation", "Random Forest / Feature Importance", "Preserves multi-variable co-variance structure required for accurate Gini impurity variance reduction."),
        ("Log Transformation", "Parametric ANOVA / t-Tests", "Normalizes right-skewed residuals to satisfy homoscedasticity and normality assumptions."),
        ("Deduplication & Timestamp Alignment", "Executive Dashboard / Time Series", "Eliminates double-counting in revenue KPIs and prevents irregular spacing in ARIMA/Prophet models.")
    ]

    for row_idx, data in enumerate(link_data, start=1):
        row_cells = table_link.rows[row_idx].cells
        row_cells[0].paragraphs[0].text = data[0]
        row_cells[1].paragraphs[0].text = data[1]
        row_cells[2].paragraphs[0].text = data[2]

    style_table_rows(table_link, widths_link)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

    # Section 8
    add_heading_1("8. Phased Execution Timeline & Project Milestones (30–35 Hours)")
    p = doc.add_paragraph(
        "The project execution is structured into five sequential phases spanning 30 to 35 labor hours."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("8.1 Visual Gantt Roadmap Chart")
    if os.path.exists("timeline_gantt.png"):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img2 = p_img2.add_run()
        run_img2.add_picture("timeline_gantt.png", width=Inches(6.2))
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap2 = p_cap2.add_run("Figure 2: 30-35 Hour Phased Preprocessing Execution Timeline")
        r_cap2.font.size = Pt(8.5)
        r_cap2.font.italic = True
        r_cap2.font.color.rgb = RGBColor(100, 116, 139)

    add_heading_2("8.2 Detailed Phase Deliverables Table")

    table_time = doc.add_table(rows=6, cols=4)
    table_time.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_time = [Inches(1.5), Inches(1.0), Inches(2.2), Inches(1.8)]
    
    headers_time = ["Project Phase", "Hours Allocated", "Core Activities", "Phase Deliverable"]
    for idx, text in enumerate(headers_time):
        table_time.rows[0].cells[idx].paragraphs[0].text = text
    style_table_header(table_time.rows[0], widths_time)

    time_data = [
        ("Phase 1: Sourcing & Setup", "6 Hours", "API credentials setup, source selection evaluation, raw data ingestion scripts.", "Data Selection Matrix & Ingestion Scripts"),
        ("Phase 2: Ingestion & Validation", "6 Hours", "Schema enforcement, data type casting, Pydantic contract verification.", "Schema Contract & Raw Staging Table"),
        ("Phase 3: Quality Audit & Cleaning", "8 Hours", "Missingness diagnostics, KNN/MICE imputation, IQR/Z-score/Isolation Forest outlier capping.", "Data Quality Report & Cleaned Dataset"),
        ("Phase 4: Feature Transformation", "8 Hours", "Robust/Standard scaling, One-Hot/Target encoding, skewness log transforms.", "Engineered Feature Store (Parquet)"),
        ("Phase 5: Integration & Docs", "7 Hours", "Pipeline modularization, unit testing, documentation, README.md and .DOC finalization.", "Final Strategy DOC & Modular Code")
    ]

    for row_idx, data in enumerate(time_data, start=1):
        row_cells = table_time.rows[row_idx].cells
        row_cells[0].paragraphs[0].text = data[0]
        row_cells[1].paragraphs[0].text = data[1]
        row_cells[2].paragraphs[0].text = data[2]
        row_cells[3].paragraphs[0].text = data[3]

    style_table_rows(table_time, widths_time)

    doc.save(output_path)
    print(f"Document successfully created and saved to {output_path}")

if __name__ == "__main__":
    create_workflow_diagram()
    create_timeline_diagram()
    generate_doc()
